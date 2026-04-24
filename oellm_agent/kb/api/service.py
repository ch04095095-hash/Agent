#!/usr/bin/env python3
"""
Knowledge Base core service.
Encapsulates all KB operations: file reading, chunking, embedding, vector search, BM25, hybrid fusion, reranking.
"""
import hashlib
import json
import pickle
import re
import shutil
import sqlite3
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

import faiss
import numpy as np
from rank_bm25 import BM25Okapi
from sentence_transformers import CrossEncoder, SentenceTransformer


# ---------------------------------------------------------------------------
# Data models
# ---------------------------------------------------------------------------

@dataclass
class KBChunk:
    chunk_id: str
    doc_id: str
    title: str
    source_type: str
    tags: List[str]
    content: str
    chunk_index: int
    source_path: str


# ---------------------------------------------------------------------------
# Text chunking
# ---------------------------------------------------------------------------

class RecursiveChunker:
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 120):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split(self, text: str) -> List[str]:
        text = re.sub(r"\n{3,}", "\n\n", text.strip())
        if not text:
            return []

        separators = ["\n\n", "\n", "。", "；", "，", " "]
        chunks = [text]

        for sep in separators:
            new_chunks: List[str] = []
            for c in chunks:
                if len(c) <= self.chunk_size:
                    new_chunks.append(c)
                else:
                    parts = c.split(sep)
                    if len(parts) == 1:
                        new_chunks.append(c)
                    else:
                        buff = ""
                        join_sep = sep if sep not in {" ", "", None} else " "
                        for p in parts:
                            p = p.strip()
                            if not p:
                                continue
                            cand = f"{buff}{join_sep if buff else ''}{p}"
                            if len(cand) <= self.chunk_size:
                                buff = cand
                            else:
                                if buff:
                                    new_chunks.append(buff)
                                buff = p
                        if buff:
                            new_chunks.append(buff)
            chunks = new_chunks

        final_chunks: List[str] = []
        for c in chunks:
            c = c.strip()
            if not c:
                continue
            if len(c) <= self.chunk_size:
                final_chunks.append(c)
            else:
                step = max(1, self.chunk_size - self.chunk_overlap)
                for i in range(0, len(c), step):
                    final_chunks.append(c[i : i + self.chunk_size])
        return final_chunks


# ---------------------------------------------------------------------------
# File reading
# ---------------------------------------------------------------------------

class FileReader:
    _ollama_base: str = "http://localhost:11434/api/generate"

    @classmethod
    def read(cls, path: Path) -> Tuple[str, Dict]:
        ext = path.suffix.lower()
        if ext in {".txt", ".md"}:
            return path.read_text(encoding="utf-8"), {}
        if ext == ".json":
            data = json.loads(path.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                return cls._json_to_text(data), {}
            return json.dumps(data, ensure_ascii=False, indent=2), {}
        if ext == ".pdf":
            return cls._read_pdf(path), {}
        if ext in {".xlsx", ".xls"}:
            return cls._read_excel(path), {}
        if ext == ".docx":
            return cls._read_docx(path), {}
        raise ValueError(f"unsupported file type: {ext}")

    @staticmethod
    def _read_pdf(path: Path) -> str:
        try:
            import pdfplumber
        except ImportError:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            pages = []
            for p in reader.pages:
                pages.append(p.extract_text() or "")
            return "\n\n".join(pages)

        text_parts = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n\n".join(text_parts)

    @staticmethod
    def _read_excel(path: Path) -> str:
        ext = path.suffix.lower()
        parts = []
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(path), data_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines = [f"[Sheet: {sheet_name}]"]
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        lines.append(" | ".join(cells))
                if len(lines) > 1:
                    parts.append("\n".join(lines))
        except ImportError:
            import xlrd
            if ext == ".xlsx":
                raise ValueError("openpyxl is required for .xlsx files")
            wb = xlrd.open_workbook(str(path))
            for sheet_idx in range(wb.nsheets):
                ws = wb.sheet_by_index(sheet_idx)
                lines = [f"[Sheet: {ws.name}]"]
                for row_idx in range(ws.nrows):
                    cells = [str(ws.cell_value(row_idx, col_idx)) for col_idx in range(ws.ncols)]
                    cells = [c for c in cells if c.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
                if len(lines) > 1:
                    parts.append("\n".join(lines))
        return "\n\n".join(parts)

    @staticmethod
    def _read_docx(path: Path) -> str:
        from docx import Document

        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)

    @staticmethod
    def _json_to_text(data: dict, indent: int = 0) -> str:
        lines = []
        for k, v in data.items():
            if isinstance(v, (dict, list)):
                lines.append(f"{'  ' * indent}{k}:")
                lines.append(FileReader._json_to_text(v, indent + 1))
            else:
                lines.append(f"{'  ' * indent}{k}: {v}")
        return "\n".join(lines)


# ---------------------------------------------------------------------------
# Embedding helper
# ---------------------------------------------------------------------------

class Embedder:
    def __init__(self, model_name: str = "BAAI/bge-small-zh-v1.5", cache_dir: Optional[str] = None):
        model_kwargs = {}
        if cache_dir:
            model_kwargs["cache_folder"] = cache_dir
        self.embedder = SentenceTransformer(model_name, **model_kwargs)
        sample_vec = self.embedder.encode(
            ["test"],
            normalize_embeddings=True,
            convert_to_numpy=True,
            show_progress_bar=False,
        )[0]
        self.vector_size = int(sample_vec.shape[0])

    def embed(self, texts: Sequence[str]) -> np.ndarray:
        arr = self.embedder.encode(
            list(texts),
            normalize_embeddings=True,
            convert_to_numpy=True,
            show_progress_bar=False,
        )
        return np.asarray(arr, dtype=np.float32)


# ---------------------------------------------------------------------------
# Reranker helper
# ---------------------------------------------------------------------------

class Reranker:
    def __init__(self, model_name: str = "BAAI/bge-reranker-base", cache_dir: Optional[str] = None):
        try:
            model_kwargs = {}
            if cache_dir:
                model_kwargs["cache_folder"] = cache_dir
            self.reranker = CrossEncoder(model_name, device="cpu", **model_kwargs)
            self._available = True
        except Exception:
            self.reranker = None
            self._available = False

    def rerank(self, query: str, candidates: List[KBChunk], top_k: int) -> List[Tuple[KBChunk, float]]:
        if not self._available or not self.reranker or not candidates:
            return [(c, 0.0) for c in candidates[:top_k]]

        pairs = [[query, c.content] for c in candidates]
        try:
            scores = self.reranker.predict(pairs)
        except Exception:
            return [(c, 0.0) for c in candidates[:top_k]]

        out = []
        for chunk, score in zip(candidates, scores):
            out.append((chunk, float(score)))
        out.sort(key=lambda x: x[1], reverse=True)
        return out[:top_k]


# ---------------------------------------------------------------------------
# Main KB service
# ---------------------------------------------------------------------------

class KBService:
    def __init__(
        self,
        db_path: Path,
        collection: str = "default",
        embed_model: str = "BAAI/bge-small-zh-v1.5",
        rerank_model: str = "BAAI/bge-reranker-base",
        embed_cache_dir: Optional[str] = None,
        rerank_cache_dir: Optional[str] = None,
        chunk_size: int = 800,
        chunk_overlap: int = 120,
        upload_dir: Path | str = "kb/uploads",
        dense_k: int = 50,
        bm25_k: int = 50,
        w_dense: float = 0.6,
        w_bm25: float = 0.4,
    ):
        self.db_path = Path(db_path)
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        self.collection = collection
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.dense_k = dense_k
        self.bm25_k = bm25_k
        self.w_dense = w_dense
        self.w_bm25 = w_bm25

        self.faiss_path = self.db_path.parent / f"{self.collection}.index"
        self.chunk_id_map_path = self.db_path.parent / f"{self.collection}.ids.pkl"

        self.embedder = Embedder(model_name=embed_model, cache_dir=embed_cache_dir)
        self.reranker = Reranker(model_name=rerank_model, cache_dir=rerank_cache_dir)
        self._chunker = RecursiveChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

        self._init_sqlite()
        self._init_faiss()

    # ------------------------------------------------------------------
    # SQLite
    # ------------------------------------------------------------------

    def _connect(self) -> sqlite3.Connection:
        conn = sqlite3.connect(str(self.db_path))
        conn.row_factory = sqlite3.Row
        return conn

    def _init_sqlite(self) -> None:
        with self._connect() as conn:
            conn.executescript(
                """
                CREATE TABLE IF NOT EXISTS kb_chunks (
                    chunk_id TEXT PRIMARY KEY,
                    doc_id TEXT NOT NULL,
                    title TEXT NOT NULL,
                    source_type TEXT NOT NULL,
                    tags_json TEXT NOT NULL,
                    content TEXT NOT NULL,
                    chunk_index INTEGER NOT NULL,
                    source_path TEXT NOT NULL,
                    created_at TEXT NOT NULL DEFAULT (datetime('now'))
                );

                CREATE INDEX IF NOT EXISTS idx_kb_source ON kb_chunks(source_type);

                CREATE TABLE IF NOT EXISTS kb_uploads (
                    doc_id TEXT PRIMARY KEY,
                    file_path TEXT NOT NULL,
                    filename TEXT NOT NULL,
                    title TEXT NOT NULL,
                    source_type TEXT NOT NULL DEFAULT 'manual',
                    tags_json TEXT NOT NULL DEFAULT '[]',
                    chunk_size INTEGER NOT NULL DEFAULT 800,
                    chunk_overlap INTEGER NOT NULL DEFAULT 120,
                    status TEXT NOT NULL DEFAULT 'pending',
                    created_at TEXT NOT NULL DEFAULT (datetime('now'))
                );
                """
            )

    # ------------------------------------------------------------------
    # FAISS
    # ------------------------------------------------------------------

    def _init_faiss(self) -> None:
        if self.faiss_path.exists():
            self.index = faiss.read_index(str(self.faiss_path))
            with open(self.chunk_id_map_path, "rb") as f:
                self.chunk_ids: List[str] = pickle.load(f)
        else:
            self.index = faiss.IndexFlatIP(self.embedder.vector_size)
            self.chunk_ids: List[str] = []

    def _save_faiss(self) -> None:
        faiss.write_index(self.index, str(self.faiss_path))
        with open(self.chunk_id_map_path, "wb") as f:
            pickle.dump(self.chunk_ids, f)

    # ------------------------------------------------------------------
    # Upload record management
    # ------------------------------------------------------------------

    def upload_file_record(
        self,
        doc_id: str,
        file_path: Path,
        filename: str,
        title: str,
        source_type: str,
        tags: List[str],
        chunk_size: int,
        chunk_overlap: int,
    ) -> None:
        with self._connect() as conn:
            conn.execute(
                """
                INSERT OR REPLACE INTO kb_uploads
                (doc_id, file_path, filename, title, source_type, tags_json, chunk_size, chunk_overlap, status)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    doc_id,
                    str(file_path),
                    filename,
                    title,
                    source_type,
                    json.dumps(tags, ensure_ascii=False),
                    chunk_size,
                    chunk_overlap,
                    "pending",
                ),
            )

    def get_upload_record(self, doc_id: str) -> Optional[sqlite3.Row]:
        with self._connect() as conn:
            row = conn.execute(
                "SELECT * FROM kb_uploads WHERE doc_id = ?", (doc_id,)
            ).fetchone()
            return row

    def update_upload_status(self, doc_id: str, status: str) -> None:
        with self._connect() as conn:
            conn.execute(
                "UPDATE kb_uploads SET status = ? WHERE doc_id = ?", (status, doc_id)
            )

    # ------------------------------------------------------------------
    # Chunking & building
    # ------------------------------------------------------------------

    @staticmethod
    def make_chunk_id(source_path: str, chunk_index: int, content: str) -> str:
        sig = hashlib.sha1(
            f"{source_path}|{chunk_index}|{content}".encode("utf-8")
        ).hexdigest()
        return sig

    def build_chunks_from_text(
        self,
        text: str,
        doc_id: str,
        title: str,
        source_type: str,
        tags: List[str],
        source_path: str,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
    ) -> List[KBChunk]:
        cs = chunk_size if chunk_size is not None else self.chunk_size
        co = chunk_overlap if chunk_overlap is not None else self.chunk_overlap
        chunker = RecursiveChunker(chunk_size=cs, chunk_overlap=co)
        pieces = chunker.split(text)
        out: List[KBChunk] = []
        for i, p in enumerate(pieces):
            chunk_id = self.make_chunk_id(source_path, i, p)
            out.append(
                KBChunk(
                    chunk_id=chunk_id,
                    doc_id=doc_id,
                    title=title,
                    source_type=source_type,
                    tags=tags,
                    content=p,
                    chunk_index=i,
                    source_path=source_path,
                )
            )
        return out

    # ------------------------------------------------------------------
    # Ingestion
    # ------------------------------------------------------------------

    def ingest_chunks(self, chunks: List[KBChunk], batch_size: int = 64) -> int:
        if not chunks:
            return 0

        with self._connect() as conn:
            for c in chunks:
                conn.execute(
                    """
                    INSERT OR REPLACE INTO kb_chunks(
                        chunk_id, doc_id, title, source_type, tags_json,
                        content, chunk_index, source_path
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        c.chunk_id,
                        c.doc_id,
                        c.title,
                        c.source_type,
                        json.dumps(c.tags, ensure_ascii=False),
                        c.content,
                        c.chunk_index,
                        c.source_path,
                    ),
                )

        for i in range(0, len(chunks), batch_size):
            batch = chunks[i : i + batch_size]
            vectors = self.embedder.embed([c.content for c in batch])
            self.index.add(vectors)
            self.chunk_ids.extend(c.chunk_id for c in batch)

        self._save_faiss()
        return len(chunks)

    def ingest_file(
        self,
        file_path: Path,
        doc_id: str,
        title: str,
        source_type: str = "manual",
        tags: Optional[List[str]] = None,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
    ) -> Tuple[int, str]:
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        dest = self.upload_dir / file_path.name
        shutil.copy2(str(file_path), str(dest))

        text, _ = FileReader.read(file_path)
        if not title:
            title = file_path.stem

        chunks = self.build_chunks_from_text(
            text=text,
            doc_id=doc_id or file_path.stem,
            title=title,
            source_type=source_type,
            tags=tags or [],
            source_path=str(dest.resolve()),
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        n = self.ingest_chunks(chunks)
        return n, doc_id or file_path.stem

    def ingest_texts(
        self,
        texts: List[str],
        doc_id: str,
        title: str,
        source_type: str = "manual",
        tags: Optional[List[str]] = [],
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
    ) -> Tuple[int, str]:
        combined = "\n\n".join(texts)
        chunks = self.build_chunks_from_text(
            text=combined,
            doc_id=doc_id,
            title=title,
            source_type=source_type,
            tags=tags or [],
            source_path=f"inline:{doc_id}",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        n = self.ingest_chunks(chunks)
        return n, doc_id

    def process_doc(
        self,
        doc_id: str,
        title: str = "",
        source_type: str = "manual",
        tags: Optional[List[str]] = None,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
    ) -> Tuple[int, str]:
        row = self.get_upload_record(doc_id)
        if not row:
            raise ValueError(f"doc_id not found: {doc_id}")
        if row["status"] == "processed":
            raise ValueError(f"doc_id already processed: {doc_id}")

        file_path = Path(row["file_path"])
        if not file_path.exists():
            raise FileNotFoundError(f"file not found: {file_path}")

        title = title or row["title"]
        source_type = source_type or row["source_type"]
        tags = tags if tags is not None else json.loads(row["tags_json"])
        chunk_size = chunk_size or row["chunk_size"]
        chunk_overlap = chunk_overlap or row["chunk_overlap"]

        text, _ = FileReader.read(file_path)
        chunks = self.build_chunks_from_text(
            text=text,
            doc_id=doc_id,
            title=title,
            source_type=source_type,
            tags=tags,
            source_path=str(file_path.resolve()),
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        n = self.ingest_chunks(chunks)
        self.update_upload_status(doc_id, "processed")
        return n, doc_id

    # ------------------------------------------------------------------
    # Tokenize (for BM25)
    # ------------------------------------------------------------------

    @staticmethod
    def _tokenize(text: str) -> List[str]:
        text = text.lower().strip()
        cjk = re.findall(r"[\u4e00-\u9fff]", text)
        cjk_bigrams = ["".join(cjk[i : i + 2]) for i in range(len(cjk) - 1)]
        words = re.findall(r"[a-z0-9_]+", text)
        return cjk_bigrams + words

    # ------------------------------------------------------------------
    # Search
    # ------------------------------------------------------------------

    def _dense_search(
        self, query: str, top_k: int, source_type: Optional[str]
    ) -> List[Dict]:
        qv = self.embedder.embed([query])[0:1]
        all_results = self.index.search(qv, min(top_k, max(1, self.index.ntotal)))

        out = []
        for idx, score in zip(all_results[1][0], all_results[0][0]):
            if idx < 0:
                continue
            chunk_id = self.chunk_ids[int(idx)]
            if source_type:
                with self._connect() as conn:
                    row = conn.execute(
                        "SELECT source_type FROM kb_chunks WHERE chunk_id = ?",
                        (chunk_id,),
                    ).fetchone()
                if not row or row["source_type"] != source_type:
                    continue
            out.append({"chunk_id": chunk_id, "dense_score": float(score)})
        return out

    def _bm25_search(
        self, query: str, top_k: int, source_type: Optional[str]
    ) -> List[Dict]:
        with self._connect() as conn:
            if source_type:
                rows = conn.execute(
                    "SELECT chunk_id, content FROM kb_chunks WHERE source_type = ?",
                    (source_type,),
                ).fetchall()
            else:
                rows = conn.execute(
                    "SELECT chunk_id, content FROM kb_chunks"
                ).fetchall()

        if not rows:
            return []

        chunk_ids = [r["chunk_id"] for r in rows]
        corpus = [r["content"] for r in rows]
        tokenized_corpus = [self._tokenize(t) for t in corpus]
        tokenized_query = self._tokenize(query)

        bm25 = BM25Okapi(tokenized_corpus)
        scores = bm25.get_scores(tokenized_query)

        idx = np.argsort(scores)[::-1][:top_k]
        out = []
        for i in idx:
            if scores[i] <= 0:
                continue
            out.append({"chunk_id": chunk_ids[int(i)], "bm25_score": float(scores[int(i)])})
        return out

    @staticmethod
    def _normalize_scores(score_map: Dict[str, float]) -> Dict[str, float]:
        if not score_map:
            return {}
        vals = np.array(list(score_map.values()), dtype=float)
        mn, mx = float(vals.min()), float(vals.max())
        if abs(mx - mn) < 1e-12:
            return {k: 1.0 for k in score_map}
        return {k: (v - mn) / (mx - mn) for k, v in score_map.items()}

    def _hybrid_fuse(
        self,
        dense_hits: List[Dict],
        bm25_hits: List[Dict],
        top_k: int,
        w_dense: float,
        w_bm25: float,
    ) -> List[Tuple[str, float, float, float]]:
        d_norm = self._normalize_scores({h["chunk_id"]: h["dense_score"] for h in dense_hits})
        b_norm = self._normalize_scores({h["chunk_id"]: h["bm25_score"] for h in bm25_hits})
        all_ids = set(d_norm) | set(b_norm)
        fused = []
        for cid in all_ids:
            ds = d_norm.get(cid, 0.0)
            bs = b_norm.get(cid, 0.0)
            score = w_dense * ds + w_bm25 * bs
            fused.append((cid, score, ds, bs))
        fused.sort(key=lambda x: x[1], reverse=True)
        return fused[:top_k]

    def search(
        self,
        query: str,
        top_k: int = 5,
        source_type: str = "",
        dense_k: int = 50,
        bm25_k: int = 50,
        w_dense: float = 0.6,
        w_bm25: float = 0.4,
        rerank: bool = True,
        rerank_top_k: int = 20,
    ) -> List[Dict]:
        st = source_type or None
        dense_hits = self._dense_search(query, dense_k, st)
        bm25_hits = self._bm25_search(query, bm25_k, st)
        fused = self._hybrid_fuse(dense_hits, bm25_hits, top_k=rerank_top_k if rerank else top_k, w_dense=w_dense, w_bm25=w_bm25)

        if not fused:
            return []

        top_ids = [x[0] for x in fused]
        qmarks = ",".join(["?"] * len(top_ids))
        with self._connect() as conn:
            rows = conn.execute(
                f"SELECT * FROM kb_chunks WHERE chunk_id IN ({qmarks})", top_ids
            ).fetchall()
        row_map = {r["chunk_id"]: r for r in rows}

        # Build KBChunk list for reranking
        candidates: List[KBChunk] = []
        for cid, score, ds, bs in fused:
            r = row_map.get(cid)
            if not r:
                continue
            candidates.append(
                KBChunk(
                    chunk_id=cid,
                    doc_id=r["doc_id"],
                    title=r["title"],
                    source_type=r["source_type"],
                    tags=json.loads(r["tags_json"]),
                    content=r["content"],
                    chunk_index=r["chunk_index"],
                    source_path=r["source_path"],
                )
            )

        # Apply reranking
        if rerank and candidates:
            reranked = self.reranker.rerank(query, candidates, top_k)
            result_list = []
            for chunk, rr_score in reranked:
                orig = next((x for x in fused if x[0] == chunk.chunk_id), None)
                ds = orig[2] if orig else 0.0
                bs = orig[3] if orig else 0.0
                result_list.append(
                    {
                        "chunk_id": chunk.chunk_id,
                        "doc_id": chunk.doc_id,
                        "title": chunk.title,
                        "source_type": chunk.source_type,
                        "tags": chunk.tags,
                        "chunk_index": chunk.chunk_index,
                        "content": chunk.content,
                        "source_path": chunk.source_path,
                        "score": rr_score,
                        "dense_score": ds,
                        "bm25_score": bs,
                        "rerank_score": rr_score,
                    }
                )
            return result_list[:top_k]

        out = []
        for cid, score, ds, bs in fused[:top_k]:
            r = row_map.get(cid)
            if not r:
                continue
            out.append(
                {
                    "chunk_id": cid,
                    "doc_id": r["doc_id"],
                    "title": r["title"],
                    "source_type": r["source_type"],
                    "tags": json.loads(r["tags_json"]),
                    "chunk_index": r["chunk_index"],
                    "content": r["content"],
                    "source_path": r["source_path"],
                    "score": score,
                    "dense_score": ds,
                    "bm25_score": bs,
                    "rerank_score": None,
                }
            )
        return out

    # ------------------------------------------------------------------
    # Stats & management
    # ------------------------------------------------------------------

    def stats(self) -> Dict:
        with self._connect() as conn:
            total = conn.execute("SELECT COUNT(*) FROM kb_chunks").fetchone()[0]
            docs = conn.execute("SELECT COUNT(DISTINCT doc_id) FROM kb_chunks").fetchone()[0]
            types = conn.execute(
                "SELECT source_type, COUNT(*) FROM kb_chunks GROUP BY source_type"
            ).fetchall()
        return {
            "total_chunks": total,
            "total_docs": docs,
            "by_source_type": dict(types),
            "indexed_vectors": self.index.ntotal,
        }

    def delete_doc(self, doc_id: str) -> int:
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT chunk_id FROM kb_chunks WHERE doc_id = ?", (doc_id,)
            ).fetchall()
            count = len(rows)
            if count == 0:
                return 0
            conn.execute("DELETE FROM kb_chunks WHERE doc_id = ?", (doc_id,))

        # Rebuild FAISS index (FAISS doesn't support in-place delete)
        self._rebuild_faiss()
        return count

    def _rebuild_faiss(self) -> None:
        with self._connect() as conn:
            rows = conn.execute("SELECT chunk_id, content FROM kb_chunks").fetchall()
        self.chunk_ids = []
        if self.index.ntotal > 0:
            self.index.reset()
        if not rows:
            self._save_faiss()
            return
        texts = [r["content"] for r in rows]
        self.chunk_ids = [r["chunk_id"] for r in rows]
        vectors = self.embedder.embed(texts)
        self.index.add(vectors)
        self._save_faiss()

    def clear(self) -> None:
        with self._connect() as conn:
            conn.execute("DELETE FROM kb_chunks")
        if self.index.ntotal > 0:
            self.index.reset()
        self.chunk_ids = []
        self._save_faiss()
